"""Build the single owner-facing GODFIN private-launch completion runbook."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "docs" / "GODFIN_OWNER_COMPLETION_RUNBOOK.docx"

NAVY = "163A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN_FILL = "EAF7EF"
GREEN = "17663A"
AMBER_FILL = "FFF4D8"
AMBER = "7A5A00"
RED_FILL = "FDECEC"
RED = "9B1C1C"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell,
    *,
    top: int = 80,
    start: int = 120,
    bottom: int = 80,
    end: int = 120,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, value, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([run_props, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Runbook Caption" not in styles:
        caption = styles.add_style("Runbook Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Runbook Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.line_spacing = 1.15


def add_numbering(document: Document, *, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(
        qn("w:val"),
        "%1." if kind == "decimal" else ("☐" if kind == "check" else "•"),
    )
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "60")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.append(p_pr)
    if kind in {"bullet", "check"}:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Segoe UI Symbol")
        fonts.set(qn("w:hAnsi"), "Segoe UI Symbol")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(document: Document, text: str, num_id: int, *, bold_prefix: str | None = None):
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    p_pr.append(num_pr)
    paragraph.paragraph_format.keep_together = True
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        prefix.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def add_callout(
    document: Document,
    label: str,
    body: str,
    *,
    tone: str = "info",
) -> None:
    fill, color = {
        "good": (GREEN_FILL, GREEN),
        "warn": (AMBER_FILL, AMBER),
        "risk": (RED_FILL, RED),
        "info": (LIGHT_BLUE, NAVY),
    }[tone]
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    borders = OxmlElement("w:pBdr")
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"), "single")
    left_border.set(qn("w:sz"), "18")
    left_border.set(qn("w:space"), "8")
    left_border.set(qn("w:color"), color)
    borders.append(left_border)
    p_pr.extend([shading, borders])
    lead = paragraph.add_run(f"{label}  ")
    set_run_font(lead, size=10.5, color=color, bold=True)
    content = paragraph.add_run(body)
    set_run_font(content, size=10.5, color=INK)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.5, color=NAVY, bold=True)
    mark_header_row(header)
    for row_values in rows:
        row = table.add_row()
        for index, text in enumerate(row_values):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, size=9.25, color=INK)
    set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_owner_fields(document: Document, fields: list[str]) -> None:
    for field in fields:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        label = paragraph.add_run(f"{field}: ")
        label.bold = True
        paragraph.add_run("_" * 74)


def add_source(document: Document, label: str, url: str) -> None:
    paragraph = document.add_paragraph(style="Runbook Caption")
    paragraph.add_run("Official reference: ")
    add_hyperlink(paragraph, label, url)


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)
    decimal_num = add_numbering(document, kind="decimal")
    bullet_num = add_numbering(document, kind="bullet")
    check_num = add_numbering(document, kind="check")
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(
        header.add_run("GODFIN  /  PRIVATE OWNER RUNBOOK"),
        size=8.5,
        color=MUTED,
        bold=True,
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run("Private & confidential  •  Page "), size=9, color=MUTED)
    add_field(footer, "PAGE")
    set_run_font(footer.add_run(" of "), size=9, color=MUTED)
    add_field(footer, "NUMPAGES")

    # Editorial-cover opening. Named override: owner_runbook_title.
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(52)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    set_run_font(kicker.add_run("PRIVATE LAUNCH OPERATIONS"), size=10, color=BLUE, bold=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(
        title.add_run("GODFIN Owner\nCompletion Runbook"),
        size=28,
        color=NAVY,
        bold=True,
    )
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_run_font(
        subtitle.add_run(
            "Owner-controlled credentials, signing, compliance, validation, and launch gates"
        ),
        size=13,
        color=DARK_BLUE,
    )
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(42)
    set_run_font(
        metadata.add_run(
            "Version 2.5  •  16 August 2026  •  oculusrex14/GODFIN (private)\n"
            "Production branch: codex/godfin-production-v5"
        ),
        size=10,
        color=MUTED,
    )
    add_callout(
        document,
        "NON-NEGOTIABLE",
        "Never paste secrets, OAuth credentials, signing certificates, real statements, local databases, PINs, license keys, payout identities, or customer financial data into this document, source control, chat, tickets, or email.",
        tone="risk",
    )
    owner = document.add_paragraph()
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    owner.paragraph_format.space_before = Pt(20)
    set_run_font(
        owner.add_run("Owner: ______________________________    Target date: __________________"),
        size=10.5,
        color=INK,
    )

    document.add_page_break()
    document.add_heading("How to use this runbook", level=1)
    for item in (
        "Read the complete runbook once. Following the owner's preferred work order, finish local code, tests, documentation, and package checks first; then complete the provider/browser steps in Sections 1–4 as one final batch. Section 16 is always last.",
        "Use provider dashboards in your own browser. Store every secret only in the provider’s encrypted environment, GitHub Actions secret store, and your business password vault.",
        "Keep the desktop backend local. Supabase, Vercel, Cashfree, Resend, and R2 are website, licensing, email, and distribution services only; they must never become the application database.",
        "Do not enable checkout, the reward pilot, update-feed promotion, or public launch content until the matching sign-off gate is complete.",
        "When a step says “ask Codex to verify,” provide only the non-secret project name, public URL, or workflow run URL—not the credential value.",
    ):
        add_list_item(document, item, decimal_num)

    document.add_page_break()
    document.add_heading("Current verified state", level=1)
    add_callout(
        document,
        "ENGINEERING BASELINE",
        "Production remediation is tracked on the private v5 branch. The clean Python 3.12 baseline is 878 passing backend tests. All 182 API operations publish an explicit success or intentional terminal status, a non-empty media-specific success schema, and the shared error contract. Frontend lint, accessibility, authentication, and production builds pass; five executable Cashfree contract tests and the website production build pass; desktop release/update tests pass 12/12; and package privacy checks pass 7/7.",
        tone="good",
    )
    for item in (
        "Private repository: https://github.com/oculusrex14/GODFIN",
        "Canonical production website: https://godfin.dev. The apex resolves to Vercel over HTTPS; https://godfin.vercel.app remains an operational fallback.",
        "Supabase project: GODFIN (ap-south-1). Local ordered migrations now run through 0006; migration 0006 adds provider-neutral Cashfree purchase/event fields and replay-safe Cashfree provisioning without deleting legacy purchase history. Remote application and pgTAP evidence remain required.",
        "Deprecated source is preserved only in private, read-only repository oculusrex14/GODFIN-OPUS46-ARCHIVE. Its 35-commit rewritten history and archival tag pass secret scanning; the obsolete local source/build workspace was moved to Trash while active Application Support data was preserved.",
        "PIN access recovery, portal-positioned calculation help, collapsible App Settings, external pricing navigation, auditable goal contributions, corrected goal simulation, recurring re-detection, atomic account routing, package privacy assertions, and the AY 2026–27 CA tax pack are implemented and tested.",
        "The website product tour uses real React application captures generated only from synthetic data, muted WebM/MP4 media with reduced-motion fallbacks, and build-time checks that prevent unreleased features from being advertised. The current public deployment is reachable at godfin.dev; the Cashfree code update must remain a preview until migration 0006 and provider tests pass.",
        "Production acceptance on 30 July 2026: website Playwright 4/4; Lighthouse performance 99, accessibility 100, SEO 100, LCP 1.73 seconds, CLS 0; required CSP, HSTS, frame, MIME, referrer, and permissions headers are present.",
        "Vercel already contains the Supabase public/server variables and LICENSE_SIGNING_SECRET. Values are encrypted and are intentionally not reproduced here.",
        "Google OAuth is active. The owner-controlled GODFIN Website project uses the rotated web client named GODFIN Website Rotated; the provider requests only openid/email/profile, the original client is revoked, and production sign-in returned successfully to /account twice on 30 July 2026.",
        "A non-revenue Max owner_test license is active for the owner account. The server stores only its hash, purchase history remains empty, and one macOS arm64 installation is verified through the normal three-device flow. The full key is retained only in macOS Keychain and encrypted local app storage.",
        "Dependency surfaces are separated into runtime, test, and frozen-build locks. The Gmail API client is an explicit runtime dependency. All four JavaScript workspaces and all three Python lock surfaces audit with no unaccepted known vulnerability; the qualified cryptography finding is documented in the signed evidence.",
        "The deterministic CycloneDX 1.6 SBOM contains 1,035 unique components with zero unresolved license identifiers. Third-party notices list all conditional licenses. Final human legal clearance is intentionally fail-closed and remains pending in supply-chain/legal-clearance.json.",
        "A fresh local macOS arm64 package was built from code commit 25fd5f7 using the locked toolchain. It starts in 3.000 seconds on first launch and 1.033 seconds on restart, preserves its database, enforces the local trust and maintenance boundaries, and remains below the 700 MB idle-memory budget at 604.5 MB across five processes. The private DMG SHA-256 is 5b3fd15ae40e9ab59ea2b344f6de9e15f37e6ad2030b1d2f330253f2ed4c98c5 and the ZIP SHA-256 is 8bd9d5b59e11e3bdbb5edfb5f882ed953a5341a178a2d77cd0a7190154ce5b91. This is an ad-hoc local test signature, not an Apple-notarized customer release.",
        "Release workflows require exact tag, commit, and package-version agreement; refuse an existing GitHub Release; publish deterministic SBOM, notices, checksums, and provenance; use immutable action SHAs; and require staged promotion plus a reviewed rollback path.",
        "Repository code evidence is current through commit aeb5cb9 on codex/godfin-production-v5. Cashfree, canonical-domain, Gmail-callback, and bounded-job changes are pushed privately; this v2.5 runbook is the next documentation-only commit.",
        "The Gmail callback trust defect is fixed in the installed app: the exact external-browser callback may reach the OAuth handler while ordinary routes still require the per-launch secret. Owner consent/sync/disconnect/reconnect must now be retried with the configured Desktop client; naraharikripa14@gmail.com is also listed as a Google test user.",
        "Final Gmail consent, Cashfree KYC/sandbox/live tests, Supabase migration 0006 deployment, Resend DNS, Google/Supabase canonical redirect updates, Apple/Windows certificates, R2, cross-platform clean-system evidence, dependency-license approval, and public-launch authorization are not yet complete.",
        "Reward pilot, sponsor card, PPP checkout, and OpenDataLoader shipping remain safely feature-gated where applicable.",
        "Nothing in this document authorizes a public release. Phase 6 starts only after explicit written public-launch authorization.",
    ):
        add_list_item(document, item, bullet_num)

    document.add_page_break()
    document.add_heading("Critical path and blockers", level=1)
    add_table(
        document,
        ["Blocker", "Why blocked", "Your intervention", "Completion evidence"],
        [
            ["Desktop Gmail OAuth", "Client setup is done and callback trust is fixed; end-to-end consent/sync has not been re-run.", "Retry Connect Gmail in the installed app and complete Google consent.", "Connect, callback, status, sync, disconnect, and reauthorization pass."],
            ["Dependency legal review", "Automated license inventory is complete; human approval is not.", "Review conditional licenses and sign legal-clearance.json without changing evidence hashes.", "Release gate reports approved and the signed record is archived."],
            ["Cashfree India", "Repository integration exists; KYC, credentials, webhook, and sandbox evidence are absent.", "Complete KYC; configure sandbox keys and all required webhook events.", "Replay-safe purchase, refund, dispute, and email flows pass."],
            ["Resend + DNS", "No sending key/domain verification.", "Verify godfin.dev and add the production key.", "SPF, DKIM, DMARC and two inbox tests pass."],
            ["Canonical redirects", "godfin.dev resolves, but Google/Supabase/Vercel settings still need owner verification.", "Make godfin.dev primary and retain the Vercel callback fallback.", "Apex, www redirect, OAuth, sitemap, and security headers pass."],
            ["Signing", "No GitHub signing secrets are configured.", "Complete Apple and Windows signing enrollment.", "Notarized/signed installers verify."],
            ["R2 updates", "No release bucket or releases.godfin.dev.", "Create R2, DNS, and least-privilege secrets.", "Immutable assets and updater metadata resolve."],
            ["Clean systems", "Only local macOS arm64 packaging is evidenced.", "Provide clean supported VMs/hardware.", "Install/upgrade/recovery matrix is signed."],
            ["Browser matrix", "Provider-backed browser checks are deferred to the final tranche.", "Run supported Chrome, Safari, Firefox, and Edge acceptance flows.", "Screenshots, traces, accessibility, and isolation evidence are archived."],
            ["Public launch", "Owner authorization has not been issued.", "Complete final gate and sign Section 16.", "Written authorization and release IDs."],
        ],
        [1500, 2100, 3000, 2760],
    )

    document.add_page_break()
    document.add_heading("1. Google OAuth through Supabase", level=1)
    add_callout(
        document,
        "BOUNDARY",
        "Google authentication is for website accounts and license management only. Do not request Gmail access here; the optional desktop Gmail integration is separate and local.",
        tone="info",
    )
    document.add_heading("1.1 Create the Google web client", level=2)
    add_callout(
        document,
        "CURRENT EVIDENCE",
        "Completed 30 July 2026. Dedicated Google Cloud project: GODFIN Website; project ID: godfin-website; project number: 173410737907. The active web client is GODFIN Website Rotated. Supabase Google authentication is enabled, the original client is revoked, and two production callback runs returned the owner account to /account. Do not reuse this website client for the optional desktop Gmail integration.",
        tone="good",
    )
    for item in (
        "Open Google Cloud Console in the business-owned project. Enable two-factor authentication for every administrator.",
        "Configure the OAuth consent screen. Use the final seller name, support email, privacy URL, terms URL, and verified domain. Choose the appropriate external/testing status for the launch stage.",
        "Request only openid, email, and profile through this website client. Do not add Gmail or finance scopes.",
        "Create an OAuth 2.0 Client ID of type Web application.",
        "Add authorized JavaScript origins: https://godfin.dev and https://www.godfin.dev. Keep https://godfin.vercel.app only as the documented operational fallback while migration is completed.",
        "In Supabase Dashboard → Authentication → Providers → Google, copy the exact callback URL. For this project it follows https://omrtkfwjauyakhvynutk.supabase.co/auth/v1/callback. Add that exact value as the Google authorized redirect URI.",
        "Copy the client ID and client secret directly from Google into the Supabase Google provider form. Do not download or commit a client-secret JSON file.",
    ):
        add_list_item(document, item, decimal_num)
    document.add_page_break()
    document.add_heading("1.2 Configure and test Supabase", level=2)
    for item in (
        "In Supabase Authentication → URL Configuration, set Site URL to https://godfin.dev.",
        "Add https://godfin.dev/auth/callback and https://www.godfin.dev/auth/callback to the redirect allow-list. Retain https://godfin.vercel.app/auth/callback as the explicit fallback.",
        "Enable the Google provider and save.",
        "Open https://godfin.dev/account in a private browser window, choose Google, complete consent, and verify return to /account.",
        "Sign out and repeat with naraharikripa14@gmail.com. Confirm neither account can see the other account’s purchases, licenses, activations, or waitlist data.",
        "Review Supabase Auth logs for redirect, consent, or email mismatch errors. Remove obsolete test clients when verification is complete.",
    ):
        add_list_item(document, item, decimal_num)
    add_source(document, "Supabase Google authentication guide", "https://supabase.com/docs/guides/auth/social-login/auth-google")
    for item in (
        "Google login succeeds from a private window.",
        "The account page shows the signed-in email and no other user’s data.",
        "Only openid/email/profile consent appears.",
        "The client secret exists only in Google/Supabase and the business vault.",
    ):
        add_list_item(document, item, check_num)
    add_callout(
        document,
        "VERIFIED",
        "The canonical origin is https://godfin.dev, the Supabase provider callback is https://omrtkfwjauyakhvynutk.supabase.co/auth/v1/callback, and the website callback is https://godfin.dev/auth/callback. The Vercel callback remains a fallback. A distinct second-account isolation test remains a pre-launch gate.",
        tone="good",
    )
    add_owner_fields(document, ["Google OAuth client name", "Test accounts used", "Completed by / date"])

    document.add_page_break()
    document.add_heading("1.3 Desktop Gmail OAuth for local ingestion", level=2)
    add_callout(
        document,
        "SEPARATE CLIENT AND SCOPE",
        "This is not the website Google login. Create a dedicated OAuth client of type Desktop app and request only https://www.googleapis.com/auth/gmail.readonly. GODFIN cannot send, edit, or delete email through this integration.",
        tone="warn",
    )
    add_callout(
        document,
        "IF YOU SAW MISSING_LAUNCH_TRUST",
        "Fully quit GODFIN, reopen /Applications/GODFIN.app, and begin a fresh Connect Gmail attempt from Settings. Do not reuse or refresh an old Google callback tab. Builds at commit 1e3f160 or later allow only the exact 127.0.0.1 Gmail callback through this boundary; every ordinary backend route still requires the active desktop launch secret.",
        tone="info",
    )
    for item in (
        "In a business-owned Google Cloud project dedicated to GODFIN Desktop Gmail, enable the Gmail API and configure the OAuth consent screen. Keep its client, scopes, and consent records separate from the GODFIN Website web client.",
        "Create an OAuth 2.0 Client ID of type Desktop app. Download the JSON and confirm its top-level object is installed and contains client_id and client_secret. Do not use a Web application JSON file.",
        "The local callback is http://127.0.0.1:5100/api/v1/auth/gmail/callback. GODFIN binds the backend to localhost by default; do not expose this callback or backend to the public internet.",
        "Move the downloaded JSON outside the repository into an owner-controlled Application Support location. Point GODFIN_GMAIL_CLIENT_SECRETS_FILE to that absolute file. For development only, backend/data/client_secret.json is supported but must remain ignored and must never be committed, attached, or shared.",
        "Alternatively set GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET in the private local process environment. Never put either value in source files, screenshots, this runbook, terminal history, or chat.",
        "Start GODFIN, open Settings, choose Connect Gmail, and review the Google consent screen. It must show read-only Gmail access and no send, modify, delete, contacts, Drive, calendar, or broad Google account scope.",
        "Complete the localhost callback. The installed build now permits only the exact external-browser GET callback to pass without the per-launch secret; host/origin validation, one-time state, PKCE, session binding, and expiry remain enforced. Verify Settings changes to connected, run the initial sync against a test mailbox, and confirm imports are assigned only to configured active account routes.",
        "Disconnect Gmail, confirm the local token is removed, then reconnect and verify reauthorization. Tokens are encrypted in local SQLite using GODFIN's stable local key; OAuth state is hashed, installation-bound, expiring, PKCE-protected, and single-use.",
        "If Google keeps the app in Testing, add each intended test mailbox as a test user and record the expiry/re-consent implications. Complete Google verification before inviting general customers if the provider requires it for the read-only scope.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Dedicated Desktop app client exists; the website client is not reused.",
        "Only gmail.readonly appears on the consent screen.",
        "The downloaded client JSON is outside the repository and protected by owner-only permissions.",
        "Connect, localhost callback, connected status, initial sync, disconnect, and reauthorization pass.",
        "No client secret or token appears in Git, logs, screenshots, chat, or the packaged application.",
    ):
        add_list_item(document, item, check_num)
    add_source(document, "Google Gmail API Python quickstart", "https://developers.google.com/workspace/gmail/api/quickstart/python")
    add_owner_fields(document, ["Google Cloud project name", "Desktop OAuth client name", "Test mailbox", "Completed by / date"])

    document.add_heading("2. Cashfree India: one-time payments only", level=1)
    add_callout(
        document,
        "HARD RULE",
        "Sell only the Pro and Max lifetime licenses. They include zero hosted AI credits. Do not create subscriptions, trials, recurring prices, credit packs, or monthly allowances.",
        tone="risk",
    )
    document.add_heading("2.1 Complete the business account", level=2)
    for item in (
        "Use a Cashfree account owned by the legal seller. Complete Payment Gateway activation, KYC, bank verification, statement descriptor, customer support details, tax settings, and two-factor authentication.",
        "Invite a backup administrator with the minimum role needed. Do not share one login.",
        "Confirm with Cashfree and your tax adviser that the account may sell one-time software licenses domestically and internationally, and how GST, invoices, chargebacks, and refunds must be handled.",
        "Enable Cashfree's available risk controls. Do not build hardware serial, persistent IP fingerprint, or payment-method fingerprint storage in GODFIN.",
    ):
        add_list_item(document, item, decimal_num)
    document.add_heading("2.2 Test and apply Supabase migration 0006", level=2)
    add_callout(
        document,
        "DO THIS BEFORE ENABLING CHECKOUT",
        "The deployed website cannot read Cashfree purchases until migration 0006 exists in Supabase. Keep CHECKOUT_ENABLED=false throughout this section. The migration is additive and preserves legacy purchase history, but a verified backup is still required before a production database change.",
        tone="warn",
    )
    for item in (
        "Open Docker Desktop and wait until it reports that its engine is running. This is used only to test an isolated local Supabase database; it does not receive the desktop app's financial data.",
        "Open Terminal and run: cd /Users/oculus/Projects/GODFIN/GODFIN_PRODUCTION/website",
        "Run: supabase start. Wait for the local services to finish starting. If the command says Docker is unavailable, fix Docker first rather than skipping the database test.",
        "Run: supabase db reset --local. Read the target carefully and continue only when it clearly says the local Supabase stack. This command is destructive to the local test database and must never be run with --linked.",
        "Run: supabase test db --local supabase/tests. Confirm all files, including 0006_cashfree_commerce.test.sql, finish successfully with no failed pgTAP assertion.",
        "Run: supabase stop. This shuts down the temporary local stack without touching the hosted project.",
        "In Supabase Dashboard, open project GODFIN → Database → Backups and verify that a recent recoverable backup exists. If the plan does not provide managed backups, make a private encrypted database backup/export using Supabase's documented method before continuing.",
        "Back in Terminal, run: supabase login. Complete the one-time browser authorization if asked; never paste the access token into this document or chat.",
        "Run: supabase link --project-ref omrtkfwjauyakhvynutk. Confirm the displayed project is GODFIN in ap-south-1.",
        "Run: supabase migration list --linked. Migration 0006 should appear locally and not yet remotely on the first run.",
        "Run: supabase db push --linked --dry-run. Read the output and confirm it proposes only the missing ordered migration 0006_cashfree_commerce.sql.",
        "Run: supabase db push --linked. Do not close Terminal until it reports success.",
        "Run: supabase migration list --linked again. Confirm 0006 appears in both the local and remote columns. Leave CHECKOUT_ENABLED=false until the sandbox matrix passes.",
    ):
        add_list_item(document, item, decimal_num)
    add_source(document, "Supabase database migrations", "https://supabase.com/docs/guides/deployment/database-migrations")

    document.add_heading("2.3 Configure sandbox variables and catalog", level=2)
    add_table(
        document,
        ["Setting", "Purpose", "Safe starting value", "Owner action"],
        [
            ["CASHFREE_CLIENT_ID", "Server authentication", "Secret; leave blank in Git", "Add sandbox value in Vercel Preview"],
            ["CASHFREE_CLIENT_SECRET", "Webhook/API HMAC secret", "Secret; leave blank in Git", "Add sandbox value in Vercel Preview"],
            ["CASHFREE_ENVIRONMENT", "Provider environment", "sandbox", "Change only after live acceptance"],
            ["CHECKOUT_ENABLED", "Global checkout kill switch", "false", "Enable only for a controlled test"],
            ["PPP_CHECKOUT_ENABLED", "Regional pricing gate", "false", "Keep off until global legal/tax review"],
            ["CASHFREE_GLOBAL_PAYMENTS_APPROVED", "Owner approval for global checkout", "false", "Set true only with documented approval"],
        ],
        [2800, 2500, 1600, 2460],
    )
    for item in (
        "The authoritative server catalog remains Pro ₹4,999 and Max ₹9,999. Both are one-time lifetime licenses and include no hosted AI credits.",
        "Add sandbox credentials to Vercel Preview first. Keep production credentials in a separate environment and business vault.",
        "Leave PPP_CHECKOUT_ENABLED=false. India remains the only enabled checkout region until Cashfree international-payment approval, billing-country evidence, and qualified legal/tax review are complete.",
        "Do not pass amount, currency, country, user ID, or entitlement from browser code. GODFIN selects and revalidates them server-side.",
        "Keep CHECKOUT_ENABLED=false until Supabase migration 0006 is applied and the complete sandbox matrix below passes.",
    ):
        add_list_item(document, item, decimal_num)
    document.add_heading("2.4 Configure the webhook", level=2)
    for item in (
        "In Cashfree's sandbox dashboard, open Developers → Webhooks (the wording may be Payment Gateway → Webhooks), choose Add webhook, and enter https://godfin.dev/api/webhook. Retain the Vercel fallback only while canonical migration is being verified.",
        "Subscribe to PAYMENT_SUCCESS, PAYMENT_FAILED, PAYMENT_USER_DROPPED, refund/auto-refund status, and dispute created/updated/closed events.",
        "Choose the latest available webhook version supported by the account. The code accepts 2026-01-01, 2025-01-01, and 2023-08-01 while using Cashfree's current 2026-01-01 order API.",
        "Cashfree signs the timestamp plus the exact raw body with HMAC-SHA256. Do not add a second parser/proxy that reformats the body before GODFIN verifies it.",
        "Use Cashfree Dashboard webhook replay to send the same paid event twice. Verify one provider event, one purchase, one license, and one email.",
    ):
        add_list_item(document, item, decimal_num)
    document.add_page_break()
    document.add_heading("2.5 Payment acceptance evidence", level=2)
    for item in (
        "Unauthenticated checkout returns a sign-in requirement.",
        "Successful Pro and Max sandbox payments provision exactly one lifetime license each.",
        "Invalid webhook signatures return HTTP 400 and create nothing.",
        "Cancelled, failed, dropped, or unpaid orders create no license.",
        "A browser return without a verified success webhook creates no license.",
        "A partial refund suspends; a full refund revokes; a merchant-won dispute restores only when no other adverse state exists.",
        "An auto-refund for a different payment attempt on the same order does not affect the paid license.",
        "The desktop fourth activation is rejected; deactivating one device allows another.",
        "A controlled live purchase is completed only after legal, tax, email, and refund procedures are ready.",
    ):
        add_list_item(document, item, check_num)
    add_source(document, "Cashfree current Payments API", "https://www.cashfree.com/docs/api-reference/payments/latest/overview")
    add_source(document, "Cashfree webhook signature verification", "https://www.cashfree.com/docs/payments/online/webhooks/signature-verification")
    add_owner_fields(document, ["Cashfree merchant ID (not keys)", "Sandbox webhook name", "Live webhook name", "Completed by / date"])

    document.add_heading("3. Resend, operational email, and DNS", level=1)
    for item in (
        "First make hello@godfin.dev a real receiving mailbox or forwarding address with your chosen email provider. Resend sends application email; it is not the support inbox. Send a message to hello@godfin.dev from an unrelated account and reply from it before continuing.",
        "Create the Resend account under the business owner, enable two-factor authentication, and add godfin.dev as the sending domain.",
        "In the authoritative DNS provider, add exactly the SPF and DKIM records supplied by Resend. Do not invent or duplicate SPF records.",
        "Add a DMARC record. Start with monitoring if necessary, review reports, then tighten with counsel/operations guidance.",
        "Create and monitor hello@godfin.dev as the single general-purpose address for license delivery, customer support, privacy/legal requests, and security reports. Use inbox labels or filters so each request type remains easy to track.",
        "Create a least-privilege production Resend API key. Store it directly as RESEND_API_KEY in Vercel Production and Preview; set RESEND_FROM_EMAIL to GODFIN <hello@godfin.dev>, NEXT_PUBLIC_SUPPORT_EMAIL to hello@godfin.dev, and NEXT_PUBLIC_PRIVACY_EMAIL to hello@godfin.dev.",
        "Redeploy. Submit the waitlist form, receive the confirmation email, confirm once, and verify a repeated confirmation does not duplicate the record.",
        "Complete a synthetic license email test to Gmail and another provider. Check SPF, DKIM, DMARC alignment, inbox placement, links, mobile layout, plain-text fallback, reply path, and resend behavior.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Resend domain shows Verified.",
        "SPF, DKIM, and DMARC checks pass.",
        "Double-opt-in waitlist works and stores no desktop financial data.",
        "License delivery works in two mailbox providers.",
        "hello@godfin.dev is monitored and its support, privacy/legal, license, and security filters are tested.",
    ):
        add_list_item(document, item, check_num)
    add_source(document, "Resend domain verification", "https://resend.com/docs/dashboard/domains/introduction")
    add_owner_fields(document, ["DNS provider", "Second test mailbox provider", "Support owner", "Completed by / date"])

    document.add_heading("4. Domain and Vercel production", level=1)
    for item in (
        "Register or confirm control of godfin.dev in a business-owned registrar account. Enable registrar lock, two-factor authentication, auto-renewal, and a recovery contact.",
        "In Vercel project godfin, add godfin.dev and www.godfin.dev. Copy only the DNS records Vercel displays into the authoritative DNS provider.",
        "Choose one canonical host. Redirect the other host permanently.",
        "Set NEXT_PUBLIC_SITE_URL=https://godfin.dev for Production after DNS and HTTPS are ready. Update Google/Supabase redirect allow-lists and the Cashfree webhook/return URLs before switching traffic.",
        "Verify HTTPS, HSTS, CSP, X-Frame-Options, nosniff, Referrer-Policy, Permissions-Policy, sitemap.xml, robots.txt, and a deliberate 404.",
        "Set immutable signed download URLs only after Section 12 clean-system validation. Do not point download buttons at unsigned local builds.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Apex and www resolve as designed over HTTPS.",
        "Account, pricing, download, privacy, terms, docs, waitlist, and callback paths pass.",
        "Google, Cashfree, and Resend use the canonical production host.",
        "No secret appears in build logs or browser JavaScript.",
    ):
        add_list_item(document, item, check_num)
    add_owner_fields(document, ["Registrar", "Renewal date", "Canonical host", "Completed by / date"])

    document.add_page_break()
    document.add_heading("5. License custody and three-device operations", level=1)
    add_callout(
        document,
        "CURRENT STATE",
        "LICENSE_SIGNING_SECRET is already present as an encrypted Vercel variable. Do not replace it casually: issued keys are derived from it. The database stores license hashes, not full license keys.",
        tone="warn",
    )
    for item in (
        "Confirm the exact current LICENSE_SIGNING_SECRET is stored in the business password vault with restricted access and a tested offline recovery copy. Never reveal the value.",
        "Record authorized custodians and an emergency access procedure.",
        "Use random installation IDs stored in OS secure storage. Do not add hardware serials, payment details, or persistent IP fingerprints.",
        "From a paid test account, activate three synthetic installations. Confirm a fourth returns ACTIVATION_LIMIT.",
        "Open /account, list the devices, deactivate one, and verify a new installation can activate.",
        "Verify offline grace, re-verification, invalid key, inactive license, lost-key resend, and local license removal. Confirm local financial data remains untouched.",
        "Before any future secret rotation, design versioned key IDs, legacy verification, reissue, rollback, and customer communication.",
        "Completed 30 July 2026: the owner license is kind owner_test, tier Max, linked to the owner’s Supabase account, absent from purchase/revenue records, stored only as a hash server-side, and activated on one macOS arm64 installation through the normal verification flow. Its full key is not reproduced in this document.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Vault and offline recovery copies are verified.",
        "Three active installations succeed; the fourth is blocked.",
        "Account-based deactivation and replacement succeed.",
        "No serial, IP fingerprint, payment method, or financial ledger is stored by licensing.",
    ):
        add_list_item(document, item, check_num)
    add_owner_fields(document, ["Vault record name (not the value)", "Authorized custodians", "Completed by / date"])

    document.add_heading("6. macOS signing and notarization", level=1)
    for item in (
        "Enroll the legal seller in the Apple Developer Program and complete identity verification.",
        "Create a Developer ID Application certificate for distribution outside the Mac App Store.",
        "Install the certificate and private key on a secured Mac, export a password-protected .p12, and store it only in the business vault.",
        "Create an Apple app-specific password for notarization and record the Team ID.",
        "Add GitHub Actions secrets MAC_CSC_LINK, MAC_CSC_KEY_PASSWORD, APPLE_ID, APPLE_APP_SPECIFIC_PASSWORD, and APPLE_TEAM_ID. Values must never appear in logs.",
        "Run the private release workflow. Verify code signature, hardened runtime, entitlements, Electron fuses, notarization, stapling, DMG, and ZIP update artifact.",
        "Install on clean macOS 13+ Apple silicon and Intel systems with Gatekeeper enabled. Do not instruct users to bypass Gatekeeper.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "codesign --verify succeeds.",
        "spctl assessment succeeds.",
        "stapler validate succeeds.",
        "Clean Apple silicon and Intel installs launch without override.",
    ):
        add_list_item(document, item, check_num)
    add_owner_fields(document, ["Apple Team ID", "Certificate expiry", "Workflow run URL", "Completed by / date"])

    document.add_page_break()
    document.add_heading("7. Windows signing", level=1)
    for item in (
        "Purchase an Authenticode certificate in the legal seller’s name from a trusted provider. Prefer the provider’s supported hardware- or cloud-protected signing workflow.",
        "Complete identity verification and document who controls signing access.",
        "If the CI path supports an exportable PFX, store the protected file in the vault and add WIN_CSC_LINK and WIN_CSC_KEY_PASSWORD to GitHub Actions secrets.",
        "Require a trusted timestamp so signatures remain valid after certificate expiry.",
        "Build the private draft Windows installer. Verify publisher, chain, timestamp, NSIS behavior, and retained app data.",
        "Install on clean Windows 10 22H2 and Windows 11 x64 systems with SmartScreen enabled. Record reputation warnings separately from signature validity.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Signature and timestamp verify.",
        "Publisher matches the legal seller.",
        "Windows 10 and 11 clean installs pass.",
        "Upgrade and uninstall preserve SQLite data unless the user explicitly deletes it.",
    ):
        add_list_item(document, item, check_num)
    add_owner_fields(document, ["Certificate provider", "Certificate expiry", "Workflow run URL", "Completed by / date"])

    document.add_page_break()
    document.add_heading("8. Cloudflare R2 and signed updates", level=1)
    for item in (
        "Create a business-owned Cloudflare account, enable two-factor authentication, and create an R2 bucket dedicated to release artifacts only.",
        "Connect releases.godfin.dev to the bucket and verify HTTPS. Do not place customer data or desktop backups in this bucket.",
        "Create a least-privilege R2 API token for release automation. Add GitHub secrets R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY, R2_ACCOUNT_ID, and R2_RELEASE_BUCKET.",
        "Keep release binaries under immutable versioned paths such as /v0.1.0/<artifact>. Never overwrite a released binary.",
        "Verify Content-Type, Content-Disposition, range requests, checksums, blockmaps, and architecture-specific latest metadata.",
        "For the first 5% update cohort, the promote-updates workflow requires PUBLISH_STAGED_RELEASE. Advancing to 25%, 50%, or 100% requires ADVANCE_AFTER_HEALTH_REVIEW. Both use the protected release-production environment and must not run before explicit public-launch authorization.",
        "Test rollback with a previously reviewed signed release using the separate rollback workflow and the exact confirmation ROLLBACK_SIGNED_RELEASE.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Anonymous read works; list/write does not.",
        "Versioned assets are immutable.",
        "SHA-256 checksums match a fresh download.",
        "Updater detects the correct OS/architecture artifact.",
        "Rollback metadata restores the reviewed previous version.",
    ):
        add_list_item(document, item, check_num)
    add_owner_fields(document, ["R2 account ID (not secret)", "Bucket", "Update domain", "Completed by / date"])

    document.add_heading("9. Market-data access for Max net worth", level=1)
    add_callout(
        document,
        "KEY CUSTODY",
        "GODFIN does not need a shared Twelve Data key. Each user supplies a key that is encrypted locally and sent only to Twelve Data when that user requests a quote.",
        tone="good",
    )
    for item in (
        "Review Twelve Data’s current plans, market coverage, attribution, caching, redistribution, and rate-limit terms with counsel or a responsible product owner.",
        "Define the supported liquid instruments and symbol format in support documentation. Do not promise universal coverage or uninterrupted quotes.",
        "Test stock, ETF, mutual-fund, crypto, bond, metal, and currency-conversion requests with a personal test key. Never place that key in Vercel, GitHub, screenshots, or fixtures.",
        "Verify missing, expired, rate-limited, invalid, and unavailable quotes show a recoverable error and never break manual assets or authoritative totals.",
        "Confirm land, property, gems, private assets, and unsupported instruments always require user-entered source, valuation date, and review/expiry date.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Terms and support scope reviewed.",
        "A test key works and remains encrypted locally.",
        "Quote failure leaves manual calculations usable.",
        "Freshness, provenance, currency, and expiry are visible.",
    ):
        add_list_item(document, item, check_num)
    add_source(document, "Twelve Data API documentation", "https://twelvedata.com/docs")
    add_owner_fields(document, ["Reviewer", "Supported-market decision date", "Completed by / date"])

    document.add_page_break()
    document.add_heading("10. Reward-pilot privacy, payouts, and controls", level=1)
    add_callout(
        document,
        "SAFE DEFAULT",
        "Keep reward_pilot disabled for launch unless every step below is complete. The ₹50,000 pilot is separate consent, off by default, and is not required for GODFIN’s core product.",
        tone="warn",
    )
    for item in (
        "Obtain written privacy and legal review of consent, participant eligibility, withdrawal, retention, cross-border processing if any, payout tax treatment, and incident response.",
        "Approve the fixed economics: ₹100 for the first accepted 90-day aggregate bundle; ₹25 per net-new verified template family up to six; ₹10 per material variant up to five; maximum ₹300 per participant and ₹50,000 overall.",
        "Fund a dedicated pilot payout budget and define who approves submissions, disputes, duplicate detection, and payouts.",
        "Deploy a dedicated HTTPS ingestion endpoint. Configure GODFIN_REWARD_PILOT_URL only in a controlled candidate build. Never accept HTTP.",
        "Keep payout identity in a separate access-controlled system from pseudonymized contributions. Define the linking token, access log, retention, and deletion process.",
        "Use the local preview/redaction screen before submission. Reject names, account/card numbers, UPI/VPA, emails, phones, addresses, exact dates, exact amounts, balances, transaction descriptions, and raw statements.",
        "Red-team the endpoint with synthetic identifiers and malformed payloads. Confirm the server rejects forbidden fields and the pilot cap cannot be exceeded under concurrency/retry.",
        "Run a small internal dry run, reconcile accepted bundles to payouts, delete test identities, and obtain written approval before enabling the feature flag.",
    ):
        add_list_item(document, item, decimal_num)
    document.add_page_break()
    for item in (
        "Legal/privacy review signed.",
        "₹50,000 cap and ₹300 participant cap are enforced server-side.",
        "Payout identity is separate from contribution data.",
        "Redaction rejection tests pass.",
        "Consent version, preview, withdrawal, retention, and audit evidence pass.",
        "Feature remains off until written enablement approval.",
    ):
        add_list_item(document, item, check_num)
    add_owner_fields(document, ["Pilot legal reviewer", "Payout owner", "Endpoint public hostname", "Approval date"])

    document.add_heading("11. OpenDataLoader benchmark decision", level=1)
    for item in (
        "Leave opendataloader_benchmark disabled in customer builds.",
        "Lawfully assemble 200–500 privacy-safe, redacted statement fixtures covering supported formats and known failure modes. Do not use customer files without specific permission.",
        "Manually label required fields and reconciliation totals. Record source format, parser outcome, manual corrections, runtime, and memory without storing financial identifiers.",
        "Provide Java 11+ only in the isolated benchmark environment; do not silently add a Java dependency to customer packaging.",
        "Run the current extractor and OpenDataLoader adapter on the same corpus. The decisive metric is complete reconciliation without manual correction, followed by field accuracy, runtime, memory, packaging size, and failure clarity.",
        "Ship only if the measured accuracy gain justifies Java packaging and all supported OS installers pass. Otherwise document the decision and retain the current extractor.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Corpus is lawfully sourced and fully redacted.",
        "Labels and reconciliation totals are independently reviewed.",
        "Both extractors run on identical fixtures.",
        "Ship/no-ship decision and evidence are archived.",
    ):
        add_list_item(document, item, check_num)
    add_owner_fields(document, ["Corpus owner", "Fixture count", "Decision", "Completed by / date"])

    document.add_heading("12. Clean-system packaging and recovery matrix", level=1)
    add_table(
        document,
        ["Platform", "Install", "Upgrade", "Offline", "Data/key preserved", "Uninstall", "Evidence"],
        [
            ["macOS 13+ arm64", "☐", "☐", "☐", "☐", "☐", ""],
            ["macOS 13+ x64", "☐", "☐", "☐", "☐", "☐", ""],
            ["Windows 10 22H2 x64", "☐", "☐", "☐", "☐", "☐", ""],
            ["Windows 11 x64", "☐", "☐", "☐", "☐", "☐", ""],
            ["Ubuntu 22.04 x64", "☐", "☐", "☐", "☐", "☐", ""],
        ],
        [1800, 750, 950, 850, 1300, 1000, 2710],
    )
    for item in (
        "Snapshot the clean system, install the signed candidate, set a synthetic PIN, import a redacted fixture, classify one item, create a backup, and record database path and app version.",
        "Restart online and offline. Confirm localhost binding, no cloud ledger, correct session behavior, encrypted secrets, and deterministic operation without AI.",
        "If Ollama is absent, confirm GODFIN opens the official installer and never silently installs. Test no-AI mode. On capable hardware, test the recommended signed-registry model path.",
        "Upgrade from the previous candidate. Confirm SQLite, encryption key, installation ID, license, settings, merchant memory, backup history, and reports remain valid.",
        "Exercise rollback and restore on a copy. Confirm pre-migration backup and restart-safe additive migrations.",
        "Uninstall using the platform default. Confirm application data is retained unless the user explicitly chooses deletion; document the manual removal path.",
        "Record cold start, common navigation p95, 10,000-row filter, 1,000-row import, and idle memory against the committed performance budgets. Reject accepted-baseline regression above 15%.",
    ):
        add_list_item(document, item, decimal_num)
    add_owner_fields(document, ["Test coordinator", "Candidate version", "Evidence folder", "Completed by / date"])

    document.add_page_break()
    document.add_heading("13. Security, privacy, payment, and recovery evidence", level=1)
    for item in (
        "Backend: 878 tests pass under clean Python 3.12, including auth/PIN, encryption, migrations, backups, merchant upsert, licenses, bounded background-job lock recovery, exact response contracts for all 182 API operations, goal ledgers and FD/RD suggestions, simulation reference vectors, recurring detection, atomic accounts, package privacy, CA tax pack, classification memory, performance, net worth, behavior insights, reward-pilot redaction, and UTC finance-fetch regressions.",
        "Frontend: lint, accessibility checks, authentication checks, and production build pass. Focused Playwright covers PIN entry, portal calculation help, goals, recurring re-detection, external pricing, settings disclosure, and CA tax-pack controls.",
        "Website: entitlement/payment contract verification, production build, real-app product chapters, reduced-motion fallbacks, mobile overflow, lifetime/no-bundled-credit pricing, security headers, checkout safe-disable behavior, and dependency audit pass. The final production URL scores Lighthouse performance 99, accessibility 100, and SEO 100, with LCP 1.73 seconds and CLS 0.",
        "Dependencies: runtime, test, and build Python locks are separated and hash-locked; frontend, website, desktop, and Playwright workspaces use npm ci. Audits report zero unaccepted known vulnerabilities. The Gmail API client is explicit in runtime requirements.",
        "Supply chain: deterministic CycloneDX 1.6 SBOM and third-party notices cover 1,035 unique components with zero unresolved license identifiers. Human review of conditional licenses remains a required fail-closed release gate.",
        "Release engineering: 12/12 desktop update/release workflow tests and 7/7 package privacy checks pass. All GitHub Actions are pinned to verified 40-character commit SHAs, and release provenance binds the exact commit, tag, package version, SBOM, notices, and checksums.",
        "Local packaging: the fresh commit-25fd5f7 macOS arm64 candidate passed strict codesign verification, first start 3.000 seconds, restart 1.033 seconds, idle memory 604.5 MB across five processes, database preservation, local trust boundary, and maintenance-boundary checks. The DMG and ZIP exact SHA-256 values are recorded in the current package evidence. Gatekeeper correctly rejects the candidate because it is not yet Apple-notarized.",
        "Secrets: gitleaks scans the complete cleaned history; no real statements, databases, tokens, keys, or customer screenshots are tracked.",
        "Recovery: empty database bootstrap, schema-revision backup, retained daily/weekly backups, restore-on-copy, upgrade, rollback, and license offline-grace tests pass.",
        "Payment: test-mode replay, amount/currency mismatch, invalid signature, unauthenticated checkout, device limit, deactivation, and resend behavior pass after provider credentials are available.",
        "Privacy: app data remains local; sponsor card is static/non-personalized; behavior insights are never used for consequential decisions; pilot remains off unless separately approved.",
    ):
        add_list_item(document, item, check_num)
    add_callout(
        document,
        "CI RELEASE GATE",
        "Do not tag a candidate until the CI run for the exact commit is green. A local pass is necessary but not sufficient.",
        tone="warn",
    )
    add_owner_fields(document, ["Release commit", "CI run URL", "Security scan evidence", "Completed by / date"])

    document.add_heading("14. Prepare the private draft release", level=1)
    for item in (
        "Choose an immutable semantic candidate tag that matches the desktop package version exactly, such as v0.1.0. The release workflow intentionally accepts only vX.Y.Z tags; do not use an -rc suffix unless the validated workflow contract is deliberately changed first.",
        "Push the exact candidate commit to the private repository and wait for every CI job to pass.",
        "Confirm all required Apple, Windows, and R2 secrets exist in GitHub Actions without revealing values.",
        "Create and push the annotated tag. The release workflow verifies exact tag/commit/version agreement, refuses an existing release for that tag, builds native macOS arm64/x64, Windows x64, and Linux x64 artifacts, verifies packages, generates SHA-256 checksums, SBOM, notices, and provenance, and creates a private draft GitHub Release.",
        "Confirm supply-chain/legal-clearance.json is approved. The workflow must remain fail-closed while legal status is pending or the recorded evidence hashes do not match.",
        "Review the draft only. Do not publish it. Download every artifact, verify checksums/signatures, and complete Section 12.",
        "Attach the final privacy-safe screenshots and private demo to internal review, not to a public release, until the owner approves.",
        "Record known limitations honestly: HDFC launch parser scope where applicable, live-market provider coverage, no recurring AI allowance, optional local/BYO AI, and external-service dependencies.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Private draft release exists and is not published.",
        "Every expected OS/architecture artifact and update manifest exists.",
        "Checksums, signatures, notarization, fuses, and clean-system evidence pass.",
        "Release notes match the shared entitlement manifest.",
        "No public updater promotion has run.",
    ):
        add_list_item(document, item, check_num)
    add_owner_fields(document, ["Candidate tag", "Draft release URL", "Release reviewer", "Completed by / date"])

    document.add_heading("15. Final website and support acceptance", level=1)
    add_callout(
        document,
        "VERIFIED ASSETS",
        "The current website tour demonstrates imports, learned classification, goals, recurring detection, CA tax packs, and local privacy using synthetic-data captures from the real React application. A muted WebM/MP4 loop, still posters, lazy loading, and prefers-reduced-motion fallbacks are included.",
        tone="good",
    )
    for item in (
        "Verify all advertised features are marked released in shared/entitlements.json and pass app acceptance tests. Hide anything unreleased.",
        "Review the final real screenshots captured from synthetic data. Confirm there are no real names, account numbers, emails, amounts from a real person, or provider keys.",
        "Play the private demo end-to-end. Verify it reflects the signed candidate and does not expose local paths, notifications, credentials, or personal data.",
        "Complete Google sign-in, double-opt-in waitlist, test/live checkout, license email, account activations, device deactivation, and download links on the canonical domain.",
        "Ensure Pro ₹4,999 and Max ₹9,999 are lifetime licenses with zero hosted credits; local AI and BYO keys consume no GODFIN credits; no hosted AI credit packs are currently sold.",
        "Review privacy, terms, refund, lifetime definition, device limit, offline grace, supported OS/banks, market-data caveats, and support contacts with qualified Indian legal/tax advisers.",
        "Assign launch-day monitoring for checkout failures, webhook retries, email delivery, downloads, license activation, support, refunds, security reports, and rollback.",
    ):
        add_list_item(document, item, decimal_num)
    for item in (
        "Website production paths and security headers pass.",
        "All sold features are implemented and released.",
        "Support and privacy mailboxes are staffed.",
        "Legal/tax/privacy approvals are archived.",
        "Signed downloads and account/device controls work.",
        "Screenshots and demo match the final candidate.",
    ):
        add_list_item(document, item, check_num)

    document.add_page_break()
    document.add_heading("16. Explicit public-launch gate", level=1)
    add_callout(
        document,
        "NO IMPLIED AUTHORIZATION",
        "A green build, deployed website, signed installer, draft release, completed runbook, or successful payment test does not authorize a public launch. Public release and Phase 6 require explicit written owner authorization.",
        tone="risk",
    )
    add_table(
        document,
        ["Gate", "Owner", "Complete", "Evidence / approval reference"],
        [
            ["Google OAuth callback", "Codex", "☒", "Production callback passed twice on 30 Jul 2026; rotated client active; original client revoked."],
            ["Second-account isolation", "", "☐", "Repeat with a distinct Google test account before public launch."],
            ["Desktop Gmail OAuth + sync", "", "☐", "Separate Desktop client; read-only consent; connect/sync/disconnect/reconnect evidence."],
            ["Dependency license clearance", "", "☐", "Human approval of conditional licenses with unchanged SBOM/notices hashes."],
            ["Cashfree + webhook + refund/tax", "", "☐", ""],
            ["Resend + DNS + support inboxes", "", "☐", ""],
            ["Canonical domain + security headers", "", "☐", "godfin.dev resolves to Vercel; provider callbacks and production deployment still require acceptance evidence."],
            ["License custody + three devices", "", "☐", ""],
            ["macOS/Windows signing", "", "☐", ""],
            ["R2 updates + rollback", "", "☐", ""],
            ["Local macOS arm64 package proof", "Codex", "☒", "Locked local candidate meets startup, memory, privacy, persistence, and trust-boundary budgets; not notarized."],
            ["Cross-platform clean-system matrix", "", "☐", "macOS x64, Windows 10/11 x64, and Ubuntu 22.04 x64 remain external gates."],
            ["Browser-family matrix", "", "☐", "Chrome, Safari, Firefox, and Edge acceptance evidence remains deferred."],
            ["Security/privacy/recovery matrix", "", "☐", ""],
            ["Legal/tax/privacy review", "", "☐", ""],
            ["Private draft + screenshots + demo", "", "☐", ""],
            ["Final public-launch authorization", "", "☐", ""],
        ],
        [3300, 1300, 1100, 3660],
    )
    add_owner_fields(
        document,
        [
            "Authorized release version",
            "Authorized canonical website",
            "Authorized public release date/time",
            "Owner’s written authorization reference",
            "Owner signature",
            "Date",
        ],
    )
    final = document.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.paragraph_format.space_before = Pt(20)
    set_run_font(
        final.add_run(
            "Only after this gate is signed may the team publish the GitHub Release, promote the update feed, enable production checkout, announce the product publicly, or begin Phase 6 growth work."
        ),
        size=11,
        color=RED,
        bold=True,
    )

    document.core_properties.title = "GODFIN Owner Completion Runbook"
    document.core_properties.subject = "Private launch readiness and owner-controlled completion gates"
    document.core_properties.author = "GODFIN"
    document.core_properties.keywords = "GODFIN, private launch, owner runbook, signing, Cashfree payments, privacy"
    return document


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(OUTPUT)
